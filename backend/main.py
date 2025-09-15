from __future__ import annotations

import os
import json
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
 
# Load environment from .env if present
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# Initialize FastAPI early so decorators below can reference it
app = FastAPI(title="Reports API", version="0.2.0")

# Allow local Streamlit frontend to call backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Fallback DB connector (ensures availability for comment/history endpoints)
def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect("reports.db")
    conn.row_factory = sqlite3.Row
    return conn

# Ensure database schema is initialized on startup
@app.on_event("startup")
def _startup_init_db() -> None:
    try:
        init_db()
    except Exception:
        # Best-effort init; errors will surface on first DB access
        pass

# Import n8n integration
try:
    from .n8n_integration import report_generator, n8n_manager, get_last_execution_id, store_execution_id
except ImportError:
    from n8n_integration import report_generator, n8n_manager, get_last_execution_id, store_execution_id

@app.get("/health")
def health() -> Dict[str, Any]:
    return {"ok": True}

@app.get("/proxy/docx")
def proxy_docx(url: str):
    """Proxy endpoint to serve DOCX files and avoid CORS issues."""
    try:
        import requests
        import io
        
        # Download the DOCX file from the S3 URL
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            # Return the file content with proper headers
            return StreamingResponse(
                io.BytesIO(response.content),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=document.docx",
                    "Access-Control-Allow-Origin": "*",
                    "Access-Control-Allow-Methods": "GET",
                    "Access-Control-Allow-Headers": "*",
                }
            )
        else:
            raise HTTPException(status_code=response.status_code, detail="Failed to fetch document")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error proxying document: {str(e)}")

@app.get("/docx/extract-text")
def extract_docx_text(url: str):
    """Extract text content from a DOCX file for editing."""
    try:
        import requests
        from docx import Document
        import io
        
        # Download the DOCX file from the S3 URL
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            # Extract text from DOCX
            docx_buffer = io.BytesIO(response.content)
            doc = Document(docx_buffer)
            
            # Extract all text content
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_content)
            
            return {
                "success": True,
                "text": full_text,
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables)
            }
        else:
            raise HTTPException(status_code=response.status_code, detail="Failed to fetch document")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting text: {str(e)}")

@app.post("/docx/save-text")
def save_docx_text(request: Dict[str, Any]):
    """Save edited text back to a DOCX file."""
    try:
        import requests
        from docx import Document
        import io
        import tempfile
        import os
        from pathlib import Path
        
        url = request.get("url")
        text = request.get("text", "")
        case_id = request.get("case_id", "unknown")
        
        if not url or not text:
            raise HTTPException(status_code=400, detail="URL and text are required")
        
        # Download the original DOCX file
        response = requests.get(url, timeout=30)
        if response.status_code != 200:
            raise HTTPException(status_code=response.status_code, detail="Failed to fetch original document")
        
        # Load the original document
        docx_buffer = io.BytesIO(response.content)
        doc = Document(docx_buffer)
        
        # Clear existing content
        for paragraph in doc.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Add the new text content
        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        # Read the modified file
        with open(tmp_path, "rb") as f:
            modified_content = f.read()
        
        # Clean up temporary file
        os.unlink(tmp_path)
        
        # Generate filename
        filename = f"{case_id}_edited_report.docx"
        
        return StreamingResponse(
            io.BytesIO(modified_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST",
                "Access-Control-Allow-Headers": "*",
            }
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving document: {str(e)}")

@app.get("/version")
def version() -> Dict[str, Any]:
    return {"version": app.version if hasattr(app, "version") else "unknown"}

@app.post("/n8n/start")
def api_n8n_start(case_id: str, username: Optional[str] = None):
    """Start the main n8n workflow and attempt to capture execution id.
    Returns JSON with { ok: bool, execution_id?: str, started?: bool, error?: str } and
    uses HTTP 202 for started, 500 for failure.
    """
    try:
        # Ensure we pass the dynamic case_id from the request; never use a default
        res = n8n_manager.trigger_main_workflow_and_capture_execution(case_id, {"case_id": case_id, "username": username})
        ok = bool(res.get("success") or res.get("started"))
        # Be lenient: even if we couldn't capture an execution id, treat trigger as accepted
        status = 202 if (ok or res.get("error")) else 500
        body: Dict[str, Any] = {
            "ok": ok,
            "execution_id": res.get("execution_id"),
            "started": bool(res.get("started")),
            "error": res.get("error"),
        }
        return JSONResponse(content=body, status_code=status)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/n8n/execution/{case_id}")
def api_n8n_get_execution(case_id: str) -> Dict[str, Any]:
    """Return the stored execution id for a case if present. If absent and API
    credentials are configured, return the latest execution id for the main workflow
    as a best-effort hint (useful for debugging).
    """
    try:
        stored = get_last_execution_id(case_id)
        if stored:
            # If we have API access, verify the stored id is still running; otherwise treat as stale
            try:
                if getattr(n8n_manager, "api_key", None) and getattr(n8n_manager, "n8n_base_url", None):
                    r0 = n8n_manager.session.get(f"{n8n_manager.n8n_base_url}/api/v1/executions/{stored}", timeout=6)
                    if r0.ok:
                        info = r0.json() or {}
                        finished = info.get("finished")
                        stopped = info.get("stoppedAt") or info.get("stopped_at")
                        if (finished is False) and (not stopped):
                            return {"execution_id": stored, "source": "stored"}
                        # else stale; fall through to latest lookup
            except Exception:
                # If verification fails, keep returning stored to avoid breaking existing behavior
                return {"execution_id": stored, "source": "stored"}
        # Best-effort hint using API
        try:
            if getattr(n8n_manager, "api_key", None) and getattr(n8n_manager, "main_workflow_id", None):
                import requests
                r = n8n_manager.session.get(f"{n8n_manager.n8n_base_url}/api/v1/executions?limit=25", timeout=6)
                if r.ok:
                    data = r.json()
                    items = data if isinstance(data, list) else (data.get("data") or [])
                    target = str(n8n_manager.main_workflow_id)
                    cand = []
                    for it in items:
                        if str(it.get("workflowId")) == target:
                            finished = it.get("finished")
                            stopped = it.get("stoppedAt") or it.get("stopped_at")
                            if (finished is False) and (not stopped):
                                cand.append(it)
                    if cand:
                        newest = sorted(cand, key=lambda it: it.get("id") or 0, reverse=True)[0]
                        return {"execution_id": newest.get("id"), "source": "latest"}
        except Exception:
            pass
        return {"execution_id": None, "source": None}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/n8n/cancel")
def api_n8n_cancel(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Cancel an n8n execution. Accepts { execution_id? , case_id? }.
    Prefers execution_id; falls back to case_id lookup. Returns { ok: bool } and details.
    """
    try:
        execution_id = (payload.get("execution_id") or "").strip()
        case_id = (payload.get("case_id") or "").strip()
        if not execution_id and not case_id:
            raise HTTPException(status_code=400, detail="execution_id or case_id required")
        # Prefer explicit execution id
        if execution_id:
            res = n8n_manager.cancel_by_execution_id(execution_id)
            return {"ok": bool(res.get("ok")), "result": res}
        # Otherwise try last execution id recorded for this case
        exec_id = get_last_execution_id(case_id)
        if exec_id:
            res = n8n_manager.cancel_by_execution_id(exec_id)
            return {"ok": bool(res.get("ok")), "result": res, "execution_id": exec_id}
        # Safety: do not perform broad cancellations when we cannot identify the execution
        raise HTTPException(status_code=409, detail="No execution_id available for case; cannot safely cancel")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/n8n/execution")
def api_n8n_execution(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Record execution id for a case. Body: { case_id, execution_id }"""
    try:
        case_id = str(payload.get("case_id") or "").strip()
        execution_id = str(payload.get("execution_id") or "").strip()
        if not case_id or not execution_id:
            # Allow clearing the stored id by sending an empty execution_id
            if case_id and not execution_id:
                store_execution_id(case_id, "")
                return {"ok": True, "cleared": True}
            raise HTTPException(status_code=400, detail="case_id and execution_id required")
        store_execution_id(case_id, execution_id)
        return {"ok": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/code-version")
def get_code_version() -> Dict[str, Any]:
    """Fetch code version from GitHub repository"""
    try:
        import requests
        import json
        import base64
        
        # GitHub API configuration
        github_token = "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"
        github_username = "samarth0211"
        repo_name = "n8n-workflows-backup"
        branch = "main"
        file_path = "state/QTgwEEZYYfbRhhPu.version"
        
        # Construct GitHub API URL
        github_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{file_path}?ref={branch}"
        
        # Make authenticated request to GitHub API
        headers = {
            "Authorization": f"token {github_token}",
            "Accept": "application/vnd.github.v3+json"
        }
        
        response = requests.get(github_url, headers=headers, timeout=10)
        if response.ok:
            data = response.json()
            if isinstance(data, dict):
                content = data.get("content")
                encoding = data.get("encoding")
                if content and encoding and encoding.lower() == "base64":
                    # Decode base64 content
                    raw_content = base64.b64decode(content).decode("utf-8", "ignore")
                    # Parse JSON content
                    version_data = json.loads(raw_content)
                    version = version_data.get("version", "—")
                    return {"code_version": version.replace(".json", "") if isinstance(version, str) else "—"}
        return {"code_version": "—", "error": f"GitHub API error: {response.status_code}"}
    except Exception as e:
        return {"code_version": "—", "error": str(e)}

@app.get("/reports/{case_id}/code-version")
def get_case_code_version(case_id: str) -> Dict[str, Any]:
    """Get stored code version for a specific case"""
    try:
        conn = sqlite3.connect("reports.db")
        cursor = conn.cursor()
        
        # Get the most recent report for this case_id
        cursor.execute("""
            SELECT code_version FROM reports 
            WHERE case_id = ? 
            ORDER BY id DESC 
            LIMIT 1
        """, (case_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if result and result[0] and result[0] != "—":
            return {"code_version": result[0]}
        else:
            return {"code_version": "—"}
            
    except Exception as e:
        return {"code_version": "—", "error": str(e)}

@app.post("/reports/{case_id}/code-version")
def update_case_code_version(case_id: str, request: Dict[str, Any]) -> Dict[str, Any]:
    """Update code version for a specific case"""
    try:
        code_version = request.get("code_version", "—")
        
        conn = sqlite3.connect("reports.db")
        cursor = conn.cursor()
        
        # Check if there are any reports for this case_id
        cursor.execute("SELECT COUNT(*) FROM reports WHERE case_id = ?", (case_id,))
        count = cursor.fetchone()[0]
        
        if count > 0:
            # Update the most recent report for this case_id
            cursor.execute("""
                UPDATE reports 
                SET code_version = ? 
                WHERE id = (SELECT id FROM reports WHERE case_id = ? ORDER BY id DESC LIMIT 1)
            """, (code_version, case_id))
        else:
            # Create a placeholder report record if none exists
            cursor.execute("""
                INSERT INTO reports (case_id, code_version, status, started_at)
                VALUES (?, ?, 'placeholder', datetime('now'))
            """, (case_id, code_version))
        
        conn.commit()
        conn.close()
        
        return {"success": True, "code_version": code_version}
        
    except Exception as e:
        return {"success": False, "error": str(e)}

# --- S3 integration (list and presign) ---
import boto3
from botocore.client import Config as BotoConfig

S3_BUCKET = os.getenv("S3_BUCKET_NAME") or os.getenv("S3_BUCKET") or "finallcpreports"
AWS_REGION = os.getenv("AWS_REGION", os.getenv("AWS_DEFAULT_REGION", "us-east-1"))

def s3_client():
    return boto3.client(
        "s3",
        region_name=AWS_REGION,
        config=BotoConfig(signature_version="s3v4"),
    )

def s3_list_versions(case_id: str) -> list[str]:
    client = s3_client()
    versions: set[str] = set()
    # Manifest first: {case_id}/Output/index.json (no List required beyond GetObject)
    try:
        man_key = f"{case_id}/Output/index.json"
        obj = client.get_object(Bucket=S3_BUCKET, Key=man_key)
        text = obj["Body"].read().decode("utf-8")
        data = json.loads(text) if text else {}
        for r in (data.get("runs") or []):
            v = (r.get("version") or "").strip()
            if v:
                versions.add(v)
    except Exception:
        pass
    # Standard layout
    prefix_std = f"reports/{case_id}/"
    paginator = client.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=prefix_std, Delimiter="/"):
        for cp in page.get("CommonPrefixes", []):
            key = cp.get("Prefix", "")
            parts = key.strip("/").split("/")
            if len(parts) >= 3:
                versions.add(parts[2])
    # Observed layout: {case_id}/Output/{YYYYMMDDHHMM}-{case}-{patient}-CompleteAIGenerated.pdf
    # Also handle new format: {YYYYMMDDHHMM}-{case_id}-CompleteAIGeneratedReport.pdf
    import re
    ai_re = re.compile(rf"^{case_id}/Output/(\d{{12}})-{case_id}-.+?-CompleteAIGenerated\\.pdf$", re.IGNORECASE)
    ai_re_new = re.compile(rf"^{case_id}/Output/(\d{{12}})-{case_id}-CompleteAIGeneratedReport\\.(pdf|docx)$", re.IGNORECASE)
    for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=f"{case_id}/Output/"):
        for obj in page.get("Contents", []):
            key = obj.get("Key", "")
            if ai_re.match(key):
                # version is timestamp-case-patient
                name = key.split("/")[-1]
                version = name.replace("-CompleteAIGenerated.pdf", "")
                versions.add(version)
            elif ai_re_new.match(key):
                # version is timestamp-case_id
                name = key.split("/")[-1]
                version = name.replace("-CompleteAIGeneratedReport.pdf", "").replace("-CompleteAIGeneratedReport.docx", "")
                versions.add(version)
    return sorted(list(versions), reverse=True)

def s3_list_cases() -> list[str]:
    client = s3_client()
    cases: list[str] = []
    paginator = client.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=S3_BUCKET, Delimiter="/"):
        for cp in page.get("CommonPrefixes", []):
            name = cp.get("Prefix", "").strip("/")
            if not name:
                continue
            # only 4-digit numeric case ids for now
            if name.isdigit() and len(name) == 4:
                cases.append(name)
    return sorted(cases)

def s3_presign(key: str, expires: int = 900) -> str:
    return s3_client().generate_presigned_url(
        "get_object",
        Params={"Bucket": S3_BUCKET, "Key": key},
        ExpiresIn=expires,
    )

def s3_presign_put(key: str, content_type: str = "application/octet-stream", expires: int = 900) -> str:
    return s3_client().generate_presigned_url(
        "put_object",
        Params={"Bucket": S3_BUCKET, "Key": key, "ContentType": content_type},
        ExpiresIn=expires,
    )

def s3_presign_post(key: str, content_type: str = "application/octet-stream", expires: int = 900) -> dict:
    return s3_client().generate_presigned_post(
        Bucket=S3_BUCKET,
        Key=key,
        Fields={"Content-Type": content_type},
        Conditions=[["content-length-range", 1, 200 * 1024 * 1024]],
        ExpiresIn=expires,
    )

# --- Simple PDF proxy to avoid CORS issues in client-side PDF.js ---
@app.get("/proxy/pdf")
def proxy_pdf(url: str):
    try:
        import requests as _req
        from urllib.parse import unquote
        target = unquote(url)
        r = _req.get(target, stream=True, timeout=20)
        if not r.ok:
            raise HTTPException(status_code=r.status_code, detail="Upstream error")
        def _iter():
            for chunk in r.iter_content(chunk_size=8192):
                if chunk:
                    yield chunk
        headers = {
            "Content-Type": "application/pdf",
            # Allow embedding in iframes from our frontend
            "X-Accel-Buffering": "no",
        }
        return StreamingResponse(_iter(), headers=headers, media_type="application/pdf")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=502, detail="Failed to fetch PDF")

# Simple download proxy with filename hint
@app.get("/proxy/download")
def proxy_download(url: str, filename: Optional[str] = None):
    try:
        import requests as _req
        from urllib.parse import unquote
        import io as _io
        target = unquote(url)
        r = _req.get(target, timeout=30)
        if not r.ok:
            raise HTTPException(status_code=r.status_code, detail="Upstream error")
        name = filename or target.split("/")[-1] or "download.bin"
        headers = {
            "Content-Disposition": f"attachment; filename=\"{name}\"",
            "Content-Type": r.headers.get("Content-Type", "application/octet-stream"),
        }
        return StreamingResponse(_io.BytesIO(r.content), headers=headers, media_type=headers["Content-Type"])
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=502, detail="Failed to download file")

# --- Upload endpoints for edited AI DOCX ---
@app.post("/s3/presign")
def api_presign_upload(body: Dict[str, Any]) -> Dict[str, Any]:
    """
    Returns a PUT presigned URL (and also a POST form alternative) for uploading
    an edited AI report DOCX to the canonical location:
      {case_id}/Output/{filename}
    Body: { case_id: str, filename: str, type: "ai" }
    """
    case_id = str(body.get("case_id") or "").strip()
    filename = str(body.get("filename") or "").strip()
    if not case_id or not filename:
        raise HTTPException(status_code=400, detail="case_id and filename required")
    # Normalize path (store directly under Output/)
    key = f"{case_id}/Output/{filename}"
    # Prefer PUT (simpler on client). Many S3 setups reject Content-Type on presigned PUT,
    # so we provide URL without enforcing Content-Type, and also include a POST form fallback.
    put_url = s3_client().generate_presigned_url(
        "put_object",
        Params={"Bucket": S3_BUCKET, "Key": key},
        ExpiresIn=900,
    )
    post = s3_presign_post(key, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", expires=900)
    return {
        "url": put_url,
        "method": "PUT",
        "key": key,
        "post": post,
    }

@app.post("/upload/ai")
def api_direct_upload_ai(case_id: str = Form(...), filename: str = Form(...), file: UploadFile = File(...)) -> Dict[str, Any]:
    """Server-side multipart upload to S3 for edited AI report DOCX."""
    if not case_id or not filename:
        raise HTTPException(status_code=400, detail="case_id and filename required")
    key = f"{case_id}/Output/{filename}"
    try:
        data = file.file.read()
        s3_client().put_object(
            Bucket=S3_BUCKET,
            Key=key,
            Body=data,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return {"ok": True, "key": key, "url": s3_presign(key)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Playwright DOCX → PDF renderer shim (uses local converters fallback) ---
@app.post("/render/docx-to-pdf")
def render_docx_to_pdf(body: Dict[str, Any]) -> Dict[str, Any]:
    """
    Accepts { url, case_id, filename } where url points to a DOCX.
    Converts to PDF (using local converters), uploads to S3 under
      {case_id}/Output/Rendered/{filename}
    and returns { url } to the PDF.
    """
    url = str(body.get("url") or "").strip()
    case_id = str(body.get("case_id") or "").strip()
    filename = str(body.get("filename") or "rendered.pdf").strip()
    if not url or not case_id:
        raise HTTPException(status_code=400, detail="url and case_id required")
    try:
        import requests as _rq, io, tempfile
        from pathlib import Path as _Path
        r = _rq.get(url, timeout=30)
        if not r.ok:
            raise HTTPException(status_code=502, detail="failed to fetch docx")
        with tempfile.TemporaryDirectory() as td:
            docx_path = _Path(td) / "input.docx"
            pdf_path = _Path(td) / "output.pdf"
            with open(docx_path, "wb") as f:
                f.write(r.content)
            if not _convert_docx_to_pdf_local(str(docx_path), str(pdf_path)):
                raise HTTPException(status_code=500, detail="conversion failed")
            key = f"{case_id}/Output/Rendered/{filename}"
            with open(pdf_path, "rb") as f:
                s3_client().put_object(Bucket=S3_BUCKET, Key=key, Body=f.read(), ContentType="application/pdf")
            return {"url": s3_presign(key)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Ensure a given object has a PDF representation (for docx inputs)
@app.get("/s3/ensure-pdf")
def api_s3_ensure_pdf(key: Optional[str] = None, url: Optional[str] = None) -> Dict[str, Any]:
    """
    Given an S3 key, if it is a DOCX, attempt to convert and return the PDF URL.
    If it is already a PDF, return the presigned PDF URL.
    """
    if not key and not url:
        raise HTTPException(status_code=400, detail="key or url required")
    client = s3_client()
    # If url provided, derive key from URL path
    if url and not key:
        try:
            from urllib.parse import urlparse, unquote
            p = urlparse(url)
            # Expecting /<key>
            key = unquote(p.path.lstrip("/"))
        except Exception:
            key = None
    if not key:
        raise HTTPException(status_code=400, detail="could not derive key from url")
    lower = key.lower()
    if lower.endswith(".pdf"):
        return {"url": s3_presign(key), "format": "pdf"}
    if lower.endswith(".docx"):
        # Reuse helper from assets flow
        try:
            # Download, convert locally, then upload sibling PDF
            import tempfile
            from pathlib import Path as _Path
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = _Path(tmpdir) / "input.docx"
                pdf_path = _Path(tmpdir) / "output.pdf"
                obj = client.get_object(Bucket=S3_BUCKET, Key=key)
                data = obj["Body"].read()
                with open(docx_path, "wb") as f:
                    f.write(data)
                if _convert_docx_to_pdf_local(str(docx_path), str(pdf_path)):
                    pdf_key = key[:-5] + ".pdf"
                    with open(pdf_path, "rb") as f:
                        client.put_object(Bucket=S3_BUCKET, Key=pdf_key, Body=f.read(), ContentType="application/pdf")
                else:
                    pdf_key = None
        except Exception:
            pdf_key = None
        if pdf_key:
            return {"url": s3_presign(pdf_key), "format": "pdf"}
        # Fallback to docx download
        return {"url": s3_presign(key), "format": "docx"}
    # Unknown extension, just presign
    return {"url": s3_presign(key), "format": "other"}

# Simple file-based cache helpers (used by older pages)
def _cache_path_for(case_id: str) -> Path:
    return ARTIFACTS_DIR / f"outputs_cache_{case_id}.json"

@app.get("/cache/cases")
def api_cache_cases() -> Dict[str, Any]:
    cases: list[str] = []
    for p in ARTIFACTS_DIR.glob("outputs_cache_*.json"):
        name = p.stem.replace("outputs_cache_", "")
        if name and name not in cases:
            cases.append(name)
    cases.sort(reverse=True)
    return {"cases": cases}

@app.get("/cache/{case_id}/outputs")
def api_cache_outputs(case_id: str) -> Dict[str, Any]:
    path = _cache_path_for(case_id)
    if path.exists():
        try:
            text = path.read_text(encoding="utf-8")
            data = json.loads(text) if text else {}
            return {"case_id": case_id, "items": data.get("items", []), "ground_truth_key": data.get("ground_truth_key")}
        except Exception:
            pass
    return {"case_id": case_id, "items": [], "ground_truth_key": None}

@app.post("/cache/{case_id}/refresh")
def api_cache_refresh(case_id: str) -> Dict[str, Any]:
    scan = api_s3_outputs(case_id)
    # Ground truth key
    client = s3_client()
    gt_key = None
    for folder in (f"{case_id}/Ground Truth/", f"{case_id}/GroundTruth/"):
        try:
            paginator = client.get_paginator("list_objects_v2")
            newest = None
            newest_time = None
            for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=folder):
                for obj in page.get("Contents", []):
                    key = obj.get("Key", "")
                    if not key or key.endswith("/"):
                        continue
                    low = key.lower()
                    if not (low.endswith(".pdf") or low.endswith(".docx")):
                        continue
                    lm = obj.get("LastModified")
                    if newest_time is None or (lm and lm > newest_time):
                        newest_time = lm
                        newest = key
            if newest:
                gt_key = newest
                if newest.lower().endswith(".pdf"):
                    break
        except Exception:
            continue
    path = _cache_path_for(case_id)
    items = [
        {"label": it.get("label"), "ai_key": it.get("ai_key"), "doctor_key": it.get("doctor_key")}
        for it in (scan.get("items") or [])
    ]
    try:
        path.write_text(json.dumps({"items": items, "ground_truth_key": gt_key}, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass
    return {"case_id": case_id, "count": len(items)}

@app.get("/cache/presign")
def api_cache_presign(key: str) -> Dict[str, Any]:
    if not key:
        raise HTTPException(status_code=400, detail="key required")
    return {"url": s3_presign(key)}

@app.get("/s3/{case_id}/versions")
def api_s3_versions(case_id: str) -> Dict[str, Any]:
    return {"case_id": case_id, "versions": s3_list_versions(case_id)}

@app.get("/s3/cases")
def api_s3_cases() -> Dict[str, Any]:
    return {"cases": s3_list_cases()}

@app.get("/s3/{case_id}/validate")
def api_s3_validate_case(case_id: str) -> Dict[str, Any]:
    """
    Validate if a case ID exists in S3 database.
    """
    try:
        # Check if case ID exists in S3
        cases = s3_list_cases()
        exists = case_id in cases
        
        if exists:
            return {
                "exists": True,
                "message": f"Case ID {case_id} found in database",
                "case_id": case_id
            }
        else:
            return {
                "exists": False,
                "message": f"Case ID {case_id} not found in database",
                "case_id": case_id,
                "available_cases": cases[:10]  # Show first 10 available cases
            }
    except Exception as e:
        return {
            "exists": False,
            "message": f"Error validating case ID: {str(e)}",
            "case_id": case_id,
            "error": str(e)
        }

@app.get("/s3/{case_id}/outputs")
def api_s3_outputs(case_id: str) -> Dict[str, Any]:
    """
    List all AI Generated outputs under {case_id}/Output/ and pair with Doctor-as-LLM when available.
    Returns presigned URLs so the frontend can render directly without extra lookups.
    """
    client = s3_client()
    prefix = f"{case_id}/Output/"
    items: list[dict[str, str]] = []
    
    # First, collect all files to avoid multiple API calls
    all_files = {}
    file_metadata = {}
    paginator = client.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=prefix):
        for obj in page.get("Contents", []):
            key = obj.get("Key", "")
            if not key or not key.lower().endswith((".pdf", ".docx")):
                continue
            name = key.split("/")[-1]
            all_files[name] = key
            file_metadata[name] = {
                "last_modified": obj.get("LastModified"),
                "size": obj.get("Size", 0)
            }

    # Now process AI generated files and match with doctor files
    for name, key in all_files.items():
        lower = name.lower()
        # Consider files that look like AI generated reports
        if "completeaigenerated" in lower or "ai_generated" in lower:
            # Attempt to derive a matching doctor report
            base = name
            for token in ("-CompleteAIGenerated.pdf", "_CompleteAIGenerated.pdf", "-CompleteAIGeneratedReport.pdf", "_CompleteAIGeneratedReport.pdf", "-AI_Generated.pdf", "_AI_Generated.pdf", "-CompleteAIGenerated.docx", "_CompleteAIGenerated.docx", "-CompleteAIGeneratedReport.docx", "_CompleteAIGeneratedReport.docx"):
                if name.endswith(token):
                    base = name[: -len(token)]
                    break

            # Candidate doctor names (try both PDF and DOCX)
            doctor_candidates = [
                f"{base}_LLM_As_Doctor.pdf",
                f"{base}-LLM_As_Doctor.pdf",
                f"{base}_LLM_As_Doctor.docx",
                f"{base}-LLM_As_Doctor.docx",
            ]
            doctor_key = None
            for dk in doctor_candidates:
                if dk in all_files:
                    doctor_key = all_files[dk]
                    break

            # Get timestamp from S3 metadata
            metadata = file_metadata.get(name, {})
            last_modified = metadata.get("last_modified")
            timestamp = None
            if last_modified:
                # Convert to UTC string format
                timestamp = last_modified.strftime("%Y-%m-%d %H:%M UTC")
            
            items.append({
                "label": name,
                "ai_url": s3_presign(key),
                "doctor_url": s3_presign(doctor_key) if doctor_key else "",
                "ai_key": key,
                "doctor_key": doctor_key or "",
                "timestamp": timestamp,
            })

    # Sort newest first by LastModified if available; fallback to name
    def _key_sort(it: dict[str, str]) -> str:
        return it.get("ai_key", it.get("label", ""))
    items = sorted(items, key=_key_sort, reverse=True)
    return {"case_id": case_id, "items": items}


# --- Metrics lookup from S3 JSON ({case_id}/Output/{version}.json) ---
@app.get("/s3/{case_id}/metrics")
def api_s3_metrics(case_id: str, version: Optional[str] = None) -> Dict[str, Any]:
    """
    Fetch metrics for a report from a JSON file stored in S3.
    Accepts versions in either order and normalizes:
      - {case_id}-{timestamp}
      - {timestamp}-{case_id}
    Example: 9999-202509110653.json for PDF 202509110653-9999-CompleteAIGeneratedReport.pdf
    """
    if not case_id:
        raise HTTPException(status_code=400, detail="case_id required")
    import re as _re, json as _json
    client = s3_client()
    candidates: list[str] = []
    # Normalize provided version or try to infer parts
    if version:
        version = version.strip()
        # If it's "timestamp-case", also try "case-timestamp"
        m = _re.match(r"^(\d{12})-(\d{3,})$", version)
        if m:
            ts, cid = m.group(1), m.group(2)
            candidates.append(f"{cid}-{ts}")
            candidates.append(f"{ts}-{cid}")
        else:
            # If it's "case-timestamp", also try reversed
            m2 = _re.match(r"^(\d{3,})-(\d{12})$", version)
            if m2:
                cid, ts = m2.group(1), m2.group(2)
                candidates.append(f"{cid}-{ts}")
                candidates.append(f"{ts}-{cid}")
            else:
                candidates.append(version)
    if not candidates:
        raise HTTPException(status_code=400, detail="version required")
    seen = set()
    ordered: list[str] = []
    for v in candidates:
        if v not in seen:
            seen.add(v)
            ordered.append(v)
    for v in ordered:
        key = f"{case_id}/Output/{v}.json"
        try:
            obj = client.get_object(Bucket=S3_BUCKET, Key=key)
            text = obj["Body"].read().decode("utf-8", "ignore")
            data = _json.loads(text) if text else {}
            # Some pipelines write a list of dicts; merge them
            if isinstance(data, list):
                merged: dict[str, any] = {}
                for item in data:
                    if isinstance(item, dict):
                        # Normalize keys with spaces as-is; keep later entries as extras
                        for k, v in item.items():
                            if k not in merged:
                                merged[k] = v
                data = merged
            if isinstance(data, dict):
                # Split known fields vs extras to preserve all data
                known_keys = {
                    "ocr_start_time", "ocrStartTime",
                    "ocr_end_time", "ocrEndTime",
                    "total_tokens_used", "totalTokensUsed",
                    "total_input_tokens", "totalInputTokens",
                    "total_output_tokens", "totalOutputTokens",
                }
                extras = {k: v for (k, v) in data.items() if k not in known_keys}
                return {
                    "ok": True,
                    "key": key,
                    "ocr_start_time": data.get("ocr_start_time") or data.get("ocrStartTime"),
                    "ocr_end_time": data.get("ocr_end_time") or data.get("ocrEndTime"),
                    "total_tokens_used": data.get("total_tokens_used") or data.get("totalTokensUsed"),
                    "total_input_tokens": data.get("total_input_tokens") or data.get("totalInputTokens"),
                    "total_output_tokens": data.get("total_output_tokens") or data.get("totalOutputTokens"),
                    "extras": extras,
                }
        except Exception:
            continue
    return {"ok": False, "error": "metrics json not found", "tried": [f"{case_id}/Output/{v}.json" for v in ordered]}


# --- Shared comments (user-agnostic) ---
class CommentIn(BaseModel):
    case_id: str
    ai_label: Optional[str] = None
    section: str
    subsection: Optional[str] = None
    username: Optional[str] = None
    severity: Optional[str] = None
    comment: str
    resolved: Optional[bool] = False


class CommentOut(BaseModel):
    id: int
    case_id: str
    ai_label: Optional[str]
    section: str
    subsection: Optional[str]
    username: Optional[str]
    severity: Optional[str]
    comment: str
    resolved: Optional[bool]
    created_at: str


@app.get("/comments/{case_id}", response_model=List[CommentOut])
def list_comments(case_id: str, ai_label: Optional[str] = None) -> List[CommentOut]:
    conn = get_conn()
    try:
        if ai_label:
            rows = conn.execute(
                "SELECT * FROM comments WHERE case_id=? AND ai_label=? ORDER BY id DESC",
                (case_id, ai_label),
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM comments WHERE case_id=? ORDER BY id DESC",
                (case_id,),
            ).fetchall()
        out: List[CommentOut] = []
        for r in rows:
            # Access by column name to be robust to legacy column order
            def _col(name: str):
                try:
                    return r[name]
                except Exception:
                    return None
            out.append(
                CommentOut(
                    id=_col("id"),
                    case_id=_col("case_id"),
                    ai_label=_col("ai_label"),
                    section=_col("section"),
                    subsection=_col("subsection"),
                    username=_col("username"),
                    severity=_col("severity"),
                    comment=_col("comment"),
                    resolved=bool(_col("resolved") or 0),
                    created_at=_col("created_at"),
                )
            )
        return out
    finally:
        conn.close()


@app.post("/comments", response_model=CommentOut)
def add_comment(payload: CommentIn) -> CommentOut:
    conn = get_conn()
    try:
        now = datetime.utcnow().isoformat() + "Z"
        conn.execute(
            """
            INSERT INTO comments (case_id, ai_label, section, subsection, username, severity, comment, resolved, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.case_id,
                (payload.ai_label or "").strip() or None,
                payload.section,
                payload.subsection,
                (payload.username or "").strip() or None,
                payload.severity,
                payload.comment,
                1 if (payload.resolved or False) else 0,
                now,
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM comments ORDER BY id DESC LIMIT 1").fetchone()
        assert row is not None
        def _c(name: str):
            try:
                return row[name]
            except Exception:
                return None
        return CommentOut(
            id=_c("id"),
            case_id=_c("case_id"),
            ai_label=_c("ai_label"),
            section=_c("section"),
            subsection=_c("subsection"),
            username=_c("username"),
            severity=_c("severity"),
            comment=_c("comment"),
            resolved=bool(_c("resolved") or 0),
            created_at=_c("created_at"),
        )
    finally:
        conn.close()


class DeleteCommentsIn(BaseModel):
    case_id: str
    ai_label: Optional[str] = None
    ids: List[int]


class ResolveCommentIn(BaseModel):
    id: int
    case_id: str
    resolved: bool = True

@app.patch("/comments/resolve", response_model=Dict[str, int])
def resolve_comment(payload: ResolveCommentIn) -> Dict[str, int]:
    conn = get_conn()
    try:
        cur = conn.execute(
            "UPDATE comments SET resolved=? WHERE id=? AND case_id=?",
            (1 if payload.resolved else 0, payload.id, payload.case_id),
        )
        conn.commit()
        return {"updated": cur.rowcount or 0}
    finally:
        conn.close()

@app.delete("/comments", response_model=Dict[str, int])
def delete_comments(payload: DeleteCommentsIn) -> Dict[str, int]:
    if not payload.ids:
        return {"deleted": 0}
    conn = get_conn()
    try:
        qmarks = ",".join(["?"] * len(payload.ids))
        params: List[Any] = list(map(int, payload.ids))
        base = f"DELETE FROM comments WHERE id IN ({qmarks}) AND case_id=?"
        params.append(payload.case_id)
        if payload.ai_label:
            base += " AND (ai_label=? OR ai_label IS NULL)"
            params.append(payload.ai_label)
        cur = conn.execute(base, tuple(params))
        conn.commit()
        return {"deleted": cur.rowcount or 0}
    finally:
        conn.close()

@app.get("/s3/{case_id}/{report_id}/assets")
def api_s3_assets(case_id: str, report_id: str) -> Dict[str, Any]:
    """
    Flexible S3 layout support.
    Supports either:
      reports/{case_id}/{report_id}/...
    or the observed bucket layout:
      {case_id}/Output/*.pdf  (generated)
      {case_id}/Ground Truth/* or GroundTruth/* (ground truth)
      {case_id}/pages/* (optional pages)
    """
    response: Dict[str, Any] = {"case_id": case_id, "report_id": report_id}
    client = s3_client()

    def newest_under(prefix: str, exts: tuple[str, ...]) -> str | None:
        paginator = client.get_paginator("list_objects_v2")
        newest_key = None
        newest_time = None
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=prefix):
            for obj in page.get("Contents", []):
                key = obj.get("Key", "")
                if not key or key.endswith("/"):
                    continue
                if not key.lower().endswith(exts):
                    continue
                lm = obj.get("LastModified")
                if newest_time is None or (lm and lm > newest_time):
                    newest_time = lm
                    newest_key = key
        return newest_key

    # Try standard layout first
    base1 = f"reports/{case_id}/{report_id}"
    gt1 = newest_under(f"{base1}/", ("ground_truth.pdf",))
    gen_html1 = newest_under(f"{base1}/", ("generated.html",))
    gen_pdf1 = newest_under(f"{base1}/", ("generated.pdf",))

    # Try observed layout (Output/ folder)
    base2 = f"{case_id}/"
    gt2 = None
    for p in (f"{base2}Ground Truth/", f"{base2}GroundTruth/"):
        gt2 = newest_under(p, (".pdf", ".docx")) or gt2
    import re
    ai_re = re.compile(rf"^{case_id}/Output/(\d{{12}})-{case_id}-.+?-CompleteAIGenerated\\.pdf$", re.IGNORECASE)
    ai_re_new = re.compile(rf"^{case_id}/Output/(\d{{12}})-{case_id}-CompleteAIGeneratedReport\\.(pdf|docx)$", re.IGNORECASE)
    doc_re = re.compile(rf"^{case_id}/Output/(\d{{12}})-{case_id}-.+?_LLM_As_Doctor\\.pdf$", re.IGNORECASE)
    doc_re_new = re.compile(rf"^{case_id}/Output/(\d{{12}})-{case_id}-LLM_As_Doctor\\.(pdf|docx)$", re.IGNORECASE)

    # Resolve AI/Doctor by requested report_id if it looks like observed version
    gen2 = None
    doc2 = None
    if report_id and report_id != "latest":
        # Try old format first
        cand_ai_old = f"{base2}Output/{report_id}-CompleteAIGenerated.pdf"
        cand_doc_old = f"{base2}Output/{report_id}_LLM_As_Doctor.pdf"
        # Try new format
        cand_ai_new_pdf = f"{base2}Output/{report_id}-CompleteAIGeneratedReport.pdf"
        cand_ai_new_docx = f"{base2}Output/{report_id}-CompleteAIGeneratedReport.docx"
        cand_doc_new_pdf = f"{base2}Output/{report_id}-LLM_As_Doctor.pdf"
        cand_doc_new_docx = f"{base2}Output/{report_id}-LLM_As_Doctor.docx"
        
        # Try to find AI report
        for cand in [cand_ai_old, cand_ai_new_pdf, cand_ai_new_docx]:
            try:
                client.head_object(Bucket=S3_BUCKET, Key=cand)
                gen2 = cand
                break
            except Exception:
                continue
        
        # Try to find Doctor report
        for cand in [cand_doc_old, cand_doc_new_pdf, cand_doc_new_docx]:
            try:
                client.head_object(Bucket=S3_BUCKET, Key=cand)
                doc2 = cand
                break
            except Exception:
                continue
    else:
        # latest: prefer manifest index.json if present
        used_manifest = False
        try:
            obj = client.get_object(Bucket=S3_BUCKET, Key=f"{base2}Output/index.json")
            text = obj["Body"].read().decode("utf-8")
            data = json.loads(text) if text else {}
            runs = data.get("runs") or []
            if runs:
                version = (runs[-1].get("version") or "").strip()
                if version:
                    cand_ai = f"{base2}Output/{version}-CompleteAIGenerated.pdf"
                    cand_doc = f"{base2}Output/{version}_LLM_As_Doctor.pdf"
                    try:
                        client.head_object(Bucket=S3_BUCKET, Key=cand_ai)
                        gen2 = cand_ai
                    except Exception:
                        gen2 = None
                    try:
                        client.head_object(Bucket=S3_BUCKET, Key=cand_doc)
                        doc2 = cand_doc
                    except Exception:
                        doc2 = None
                    used_manifest = True
        except Exception:
            used_manifest = False

        if not used_manifest:
            # Fallback: scan Output and pick newest AI and matching Doctor
            newest_ai_key = None
            newest_ai_time = None
            paginator = client.get_paginator("list_objects_v2")
            for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=f"{base2}Output/"):
                for obj in page.get("Contents", []):
                    key = obj.get("Key", "")
                    if not key or not key.lower().endswith((".pdf", ".docx")):
                        continue
                    if not (ai_re.match(key) or ai_re_new.match(key)):
                        continue
                    lm = obj.get("LastModified")
                    if newest_ai_time is None or (lm and lm > newest_ai_time):
                            newest_ai_time = lm
                            newest_ai_key = key
                if newest_ai_key:
                    gen2 = newest_ai_key
                    # Build doctor key from the same prefix - try both old and new formats
                    name = newest_ai_key.split("/")[-1]
                    if ai_re.match(newest_ai_key):
                        # Old format
                        version = name.replace("-CompleteAIGenerated.pdf", "")
                        cand_doc = f"{base2}Output/{version}_LLM_As_Doctor.pdf"
                    else:
                        # New format
                        version = name.replace("-CompleteAIGeneratedReport.pdf", "").replace("-CompleteAIGeneratedReport.docx", "")
                        cand_doc = f"{base2}Output/{version}-LLM_As_Doctor.pdf"
                    
                    # Try both PDF and DOCX for doctor report
                    for ext in [".pdf", ".docx"]:
                        try_cand = cand_doc.replace(".pdf", ext)
                        try:
                            client.head_object(Bucket=S3_BUCKET, Key=try_cand)
                            doc2 = try_cand
                            break
                        except Exception:
                            continue

        # Prefer standard if present, else observed
        if gt1:
            response["ground_truth_pdf"] = s3_presign(gt1)
        elif gt2:
            # Might be docx; the UI can offer download, not inline view
            response["ground_truth"] = s3_presign(gt2)
        if gen_html1:
            response["generated_html"] = s3_presign(gen_html1)
        if gen_pdf1:
            response["generated_pdf"] = s3_presign(gen_pdf1)
        if not gen_html1 and not gen_pdf1 and gen2:
            if gen2.lower().endswith(".html"):
                response["generated_html"] = s3_presign(gen2)
            else:
                response["generated_pdf"] = s3_presign(gen2)
        if doc2:
            response["doctor_pdf"] = s3_presign(doc2)

        # Optional comparison under standard layout
        comps: list[str] = []
        comp_prefix = f"{base1}/comparison/"
        paginator = client.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=comp_prefix):
            for obj in page.get("Contents", []):
                key = obj.get("Key", "")
                if not key or key.endswith("/"):
                    continue
                name = key.split("/")[-1]
                version = name.rsplit(".", 1)[0]
                if version not in comps:
                    comps.append(version)
        response["comparison_versions"] = sorted(comps, reverse=True)
        return response

    # Alias for latest assets to match frontend contract
    @app.get("/s3/{case_id}/latest/assets")
    def api_s3_latest_assets(case_id: str) -> Dict[str, Any]:
        return api_s3_assets(case_id, "latest")

    @app.get("/s3/{case_id}/{report_id}/comparison/{version}")
    def api_s3_comparison(case_id: str, report_id: str, version: str) -> Dict[str, Any]:
        base = f"reports/{case_id}/{report_id}/comparison/{version}"
        client = s3_client()
        # Prefer html then pdf
        for ext in ("html", "pdf"):
            key = f"{base}.{ext}"
            try:
                client.head_object(Bucket=S3_BUCKET, Key=key)
                return {"url": s3_presign(key), "format": ext}
            except Exception:
                continue
        raise HTTPException(status_code=404, detail="Comparison version not found")


    DB_PATH = Path(os.getenv("REPORTS_DB", "reports.db")).resolve()
    ARTIFACTS_DIR = Path(os.getenv("ARTIFACTS_DIR", "artifacts")).resolve()
    ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)


    # --- Pydantic models (defined BEFORE route decorators) ---
    class ReportIn(BaseModel):
        case_id: str
        email: Optional[str] = None
        status: Optional[str] = "queued"
        started_at: Optional[str] = None
        finished_at: Optional[str] = None
        s3_key: Optional[str] = None
        file_path: Optional[str] = None
        file_size: Optional[int] = None
        checksum: Optional[str] = None
        metadata: Optional[Dict[str, Any]] = None


    class ReportOut(BaseModel):
        id: int
        case_id: str
        email: Optional[str]
        status: Optional[str]
        started_at: Optional[str]
        finished_at: Optional[str]
        s3_key: Optional[str]
        file_path: Optional[str]
        file_size: Optional[int]
        checksum: Optional[str]
        metadata: Optional[Dict[str, Any]]


    class UserIn(BaseModel):
        username: str
        email: Optional[str] = None
        full_name: Optional[str] = None


    class UserOut(BaseModel):
        id: int
        username: str
        email: Optional[str]
        full_name: Optional[str]
        created_at: str


    class CycleIn(BaseModel):
        username: Optional[str] = None
        user_id: Optional[int] = None
        case_id: str
        status: Optional[str] = "processing"
        metadata: Optional[Dict[str, Any]] = None


    class CycleOut(BaseModel):
        id: int
        user_id: int
        case_id: str
        status: Optional[str]
        started_at: Optional[str]
        finished_at: Optional[str]
        metadata: Optional[Dict[str, Any]]


    class FileOut(BaseModel):
        id: int
        cycle_id: int
        kind: Optional[str]
        file_name: Optional[str]
        file_path: Optional[str]
        file_size: Optional[int]
        checksum: Optional[str]
        created_at: Optional[str]


    def get_conn() -> sqlite3.Connection:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        return conn


    def init_db() -> None:
        conn = get_conn()
        try:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                case_id TEXT NOT NULL,
                email TEXT,
                status TEXT,
                started_at TEXT,
                finished_at TEXT,
                s3_key TEXT,
                file_path TEXT,
                file_size INTEGER,
                checksum TEXT,
                metadata TEXT,
                code_version TEXT
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_reports_case ON reports(case_id)")
            
            # Migration: Add code_version column if it doesn't exist
            try:
                conn.execute("ALTER TABLE reports ADD COLUMN code_version TEXT")
            except sqlite3.OperationalError:
                # Column already exists, ignore
                pass
            
            # New normalized schema
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE,
                email TEXT,
                full_name TEXT,
                created_at TEXT
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_users_username ON users(username)")

            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS report_cycles (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                case_id TEXT NOT NULL,
                status TEXT,
                started_at TEXT,
                finished_at TEXT,
                metadata TEXT,
                FOREIGN KEY(user_id) REFERENCES users(id)
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_cycles_user ON report_cycles(user_id)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_cycles_case ON report_cycles(case_id)")

            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS report_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                cycle_id INTEGER NOT NULL,
                kind TEXT,
                file_name TEXT,
                file_path TEXT,
                file_size INTEGER,
                checksum TEXT,
                created_at TEXT,
                FOREIGN KEY(cycle_id) REFERENCES report_cycles(id)
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_files_cycle ON report_files(cycle_id)")
            # Shared comments across users
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS comments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                case_id TEXT NOT NULL,
                ai_label TEXT,
                section TEXT,
                subsection TEXT,
                username TEXT,
                severity TEXT,
                comment TEXT,
                resolved INTEGER DEFAULT 0,
                created_at TEXT
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_comments_case ON comments(case_id)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_comments_ai_label ON comments(ai_label)")

            # Migration: Add subsection column if it doesn't exist
            try:
                cols = [r[1] for r in conn.execute("PRAGMA table_info(comments)").fetchall()]
                if "subsection" not in cols:
                    conn.execute("ALTER TABLE comments ADD COLUMN subsection TEXT")
                if "username" not in cols:
                    conn.execute("ALTER TABLE comments ADD COLUMN username TEXT")
                if "resolved" not in cols:
                    conn.execute("ALTER TABLE comments ADD COLUMN resolved INTEGER DEFAULT 0")
                if "ai_page" in cols and "gt_page" in cols:
                    # Keep old columns for backward compatibility but they're deprecated
                    pass
            except Exception:
                pass

            # History of outputs per user+case (S3 keys only; URLs are presigned on demand)
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS outputs_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                case_id TEXT NOT NULL,
                label TEXT,
                ai_key TEXT,
                doctor_key TEXT,
                ground_truth_key TEXT,
                created_at TEXT
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_hist_case ON outputs_history(case_id)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_hist_user_case ON outputs_history(user_id, case_id)")
            # Lightweight migration: add user_id to existing DBs if missing
            try:
                cols = [r[1] for r in conn.execute("PRAGMA table_info(outputs_history)").fetchall()]
                if "user_id" not in cols:
                    conn.execute("ALTER TABLE outputs_history ADD COLUMN user_id INTEGER")
                    conn.execute("CREATE INDEX IF NOT EXISTS idx_hist_user_case ON outputs_history(user_id, case_id)")
            except Exception:
                pass

            # Runs table for webhook finalization payloads (stores latest artifacts/metrics)
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS runs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                case_id TEXT NOT NULL,
                created_at TEXT,
                ai_url TEXT,
                doc_url TEXT,
                pdf_url TEXT,
                ocr_start_time TEXT,
                ocr_end_time TEXT,
                total_tokens_used INTEGER,
                total_input_tokens INTEGER,
                total_output_tokens INTEGER
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_runs_case ON runs(case_id)")
            
            # Progress updates table for real-time progress tracking
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS progress_updates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                case_id TEXT NOT NULL,
                progress INTEGER NOT NULL,
                step INTEGER NOT NULL,
                message TEXT,
                timestamp TEXT,
                created_at TEXT DEFAULT (datetime('now'))
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_progress_case ON progress_updates(case_id)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_progress_timestamp ON progress_updates(timestamp)")
            conn.commit()
        finally:
            conn.close()


    # --- Users ---
    @app.post("/users", response_model=UserOut)
    def upsert_user(user: UserIn) -> UserOut:
        conn = get_conn()
        try:
            now = datetime.utcnow().isoformat()
            existing = conn.execute("SELECT * FROM users WHERE username=?", (user.username,)).fetchone()
            if existing:
                conn.execute(
                    "UPDATE users SET email=COALESCE(?, email), full_name=COALESCE(?, full_name) WHERE username=?",
                    (user.email, user.full_name, user.username),
                )
            else:
                conn.execute(
                    "INSERT INTO users (username, email, full_name, created_at) VALUES (?, ?, ?, ?)",
                    (user.username, user.email, user.full_name, now),
                )
            conn.commit()
            row = conn.execute("SELECT * FROM users WHERE username=?", (user.username,)).fetchone()
            assert row is not None
            return UserOut(
                id=row["id"],
                username=row["username"],
                email=row["email"],
                full_name=row["full_name"],
                created_at=row["created_at"],
            )
        finally:
            conn.close()


    @app.get("/users/{username}", response_model=UserOut)
    def get_user(username: str) -> UserOut:
        conn = get_conn()
        try:
            row = conn.execute("SELECT * FROM users WHERE username=?", (username,)).fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="User not found")
            return UserOut(
                id=row["id"],
                username=row["username"],
                email=row["email"],
                full_name=row["full_name"],
                created_at=row["created_at"],
            )
        finally:
            conn.close()


    def _ensure_user(conn: sqlite3.Connection, username: Optional[str], user_id: Optional[int], email: Optional[str], full_name: Optional[str]) -> int:
        """Return a user id for the given identity, creating a user if needed."""
        if user_id:
            row = conn.execute("SELECT id FROM users WHERE id=?", (user_id,)).fetchone()
            if row:
                return int(row[0])
        uname = (username or "anonymous").strip().lower()
        row = conn.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone()
        if row:
            return int(row[0])
        now = datetime.utcnow().isoformat()
        conn.execute(
            "INSERT INTO users (username, email, full_name, created_at) VALUES (?, ?, ?, ?)",
            (uname, email, full_name, now),
        )
        conn.commit()
        row2 = conn.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone()
        assert row2 is not None
        return int(row2[0])


    # --- Report cycles ---
    @app.post("/cycles", response_model=CycleOut)
    def create_cycle(cycle: CycleIn) -> CycleOut:
        conn = get_conn()
        try:
            user_id = _ensure_user(conn, cycle.username, cycle.user_id, None, None)
            now = datetime.utcnow().isoformat()
            conn.execute(
                """
                INSERT INTO report_cycles (user_id, case_id, status, started_at, metadata)
                VALUES (?, ?, ?, ?, ?)
                """,
                (user_id, cycle.case_id, cycle.status, now, json.dumps(cycle.metadata or {})),
            )
            conn.commit()
            row = conn.execute(
                "SELECT * FROM report_cycles WHERE user_id=? AND case_id=? ORDER BY id DESC LIMIT 1",
                (user_id, cycle.case_id),
            ).fetchone()
            assert row is not None
            return CycleOut(
                id=row["id"],
                user_id=row["user_id"],
                case_id=row["case_id"],
                status=row["status"],
                started_at=row["started_at"],
                finished_at=row["finished_at"],
                metadata=json.loads(row["metadata"]) if row["metadata"] else None,
            )
        finally:
            conn.close()


    @app.get("/users/{username}/cycles", response_model=List[CycleOut])
    def list_cycles(username: str, limit: int = 20) -> List[CycleOut]:
        conn = get_conn()
        try:
            user_row = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
            if not user_row:
                return []
            uid = int(user_row[0])
            rows = conn.execute(
                "SELECT * FROM report_cycles WHERE user_id=? ORDER BY id DESC LIMIT ?",
                (uid, limit * 2),
            ).fetchall()
            result: List[CycleOut] = []
            for r in rows:
                meta = json.loads(r["metadata"]) if r["metadata"] else None
                # Skip demo/seeded cycles
                if isinstance(meta, dict) and meta.get("source") == "seed":
                    continue
                result.append(
                    CycleOut(
                        id=r["id"],
                        user_id=r["user_id"],
                        case_id=r["case_id"],
                        status=r["status"],
                        started_at=r["started_at"],
                        finished_at=r["finished_at"],
                        metadata=meta,
                    )
                )
                if len(result) >= limit:
                    break
            return result
        finally:
            conn.close()


    # --- Files ---
    @app.get("/cycles/{cycle_id}/files", response_model=List[FileOut])
    def list_files(cycle_id: int) -> List[FileOut]:
        conn = get_conn()
        try:
            rows = conn.execute("SELECT * FROM report_files WHERE cycle_id=? ORDER BY id", (cycle_id,)).fetchall()
            out: List[FileOut] = []
            for r in rows:
                out.append(
                    FileOut(
                        id=r["id"],
                        cycle_id=r["cycle_id"],
                        kind=r["kind"],
                        file_name=r["file_name"],
                        file_path=r["file_path"],
                        file_size=r["file_size"],
                        checksum=r["checksum"],
                        created_at=r["created_at"],
                    )
                )
            return out
        finally:
            conn.close()


    # --- Cycle details & updates ---
    @app.get("/cycles/{cycle_id}", response_model=CycleOut)
    def get_cycle(cycle_id: int) -> CycleOut:
        conn = get_conn()
        try:
            row = conn.execute("SELECT * FROM report_cycles WHERE id=?", (cycle_id,)).fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Cycle not found")
            return CycleOut(
                id=row["id"],
                user_id=row["user_id"],
                case_id=row["case_id"],
                status=row["status"],
                started_at=row["started_at"],
                finished_at=row["finished_at"],
                metadata=json.loads(row["metadata"]) if row["metadata"] else None,
            )
        finally:
            conn.close()


    class CyclePatch(BaseModel):
        status: Optional[str] = None
        finished_at: Optional[str] = None
        accuracy_score: Optional[float] = None
        key_differences: Optional[int] = None
        processing_seconds: Optional[int] = None
        metadata: Optional[Dict[str, Any]] = None


    @app.patch("/cycles/{cycle_id}", response_model=CycleOut)
    def update_cycle(cycle_id: int, patch: CyclePatch) -> CycleOut:
        conn = get_conn()
        try:
            row = conn.execute("SELECT * FROM report_cycles WHERE id=?", (cycle_id,)).fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Cycle not found")
            # merge metadata
            existing_meta: Dict[str, Any] = {}
            if row["metadata"]:
                try:
                    existing_meta = json.loads(row["metadata"]) or {}
                except Exception:
                    existing_meta = {}
            # attach metrics if provided
            if patch.accuracy_score is not None:
                existing_meta["accuracy_score"] = patch.accuracy_score
            if patch.key_differences is not None:
                existing_meta["key_differences"] = patch.key_differences
            if patch.processing_seconds is not None:
                existing_meta["processing_seconds"] = patch.processing_seconds
            if patch.metadata:
                existing_meta.update(patch.metadata)

            conn.execute(
                "UPDATE report_cycles SET status=COALESCE(?, status), finished_at=COALESCE(?, finished_at), metadata=? WHERE id=?",
                (patch.status, patch.finished_at, json.dumps(existing_meta), cycle_id),
            )
            conn.commit()
            return get_cycle(cycle_id)
        finally:
            conn.close()


    @app.post("/cycles/{cycle_id}/files", response_model=FileOut)
    def upload_cycle_file(
        cycle_id: int,
        file: UploadFile = File(...),
        kind: Optional[str] = Form(None),
    ) -> FileOut:
        suffix = Path(file.filename or "artifact.pdf").suffix
        safe_name = f"cycle{cycle_id}_{int(datetime.utcnow().timestamp())}{suffix}"
        target = ARTIFACTS_DIR / safe_name
        with target.open("wb") as out:
            out.write(file.file.read())
        size = target.stat().st_size

        conn = get_conn()
        try:
            now = datetime.utcnow().isoformat()
            conn.execute(
                """
                INSERT INTO report_files (cycle_id, kind, file_name, file_path, file_size, checksum, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (cycle_id, kind, file.filename, str(target), size, None, now),
            )
            conn.commit()
            row = conn.execute(
                "SELECT * FROM report_files WHERE cycle_id=? ORDER BY id DESC LIMIT 1",
                (cycle_id,),
            ).fetchone()
            assert row is not None
            return FileOut(
                id=row["id"],
                cycle_id=row["cycle_id"],
                kind=row["kind"],
                file_name=row["file_name"],
                file_path=row["file_path"],
                file_size=row["file_size"],
                checksum=row["checksum"],
                created_at=row["created_at"],
            )
        finally:
            conn.close()


    @app.get("/files/{file_id}/download")
    def download_file(file_id: int):
        conn = get_conn()
        try:
            row = conn.execute("SELECT * FROM report_files WHERE id=?", (file_id,)).fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="File not found")
            path = row["file_path"]
            if not path or not Path(path).exists():
                raise HTTPException(status_code=404, detail="File missing on server")
            filename = Path(path).name
            # Serve PDFs inline so they can render in iframe on the Results page
            if filename.lower().endswith(".pdf"):
                return FileResponse(
                    path=path,
                    filename=filename,
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"inline; filename=\"{filename}\""},
                )
            return FileResponse(path=path, filename=filename)
        finally:
            conn.close()


    @app.post("/history/{case_id}/sync")
    def history_sync(case_id: str, username: Optional[str] = None, user_id: Optional[int] = None) -> Dict[str, Any]:
        scan = api_s3_outputs(case_id)
        # determine ground truth key
        client = s3_client()
        gt_key = None
        for folder in (f"{case_id}/Ground Truth/", f"{case_id}/GroundTruth/"):
            try:
                paginator = client.get_paginator("list_objects_v2")
                newest = None
                newest_time = None
                for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=folder):
                    for obj in page.get("Contents", []):
                        key = obj.get("Key", "")
                        if not key or key.endswith("/"):
                            continue
                        low = key.lower()
                        if not (low.endswith(".pdf") or low.endswith(".docx")):
                            continue
                        lm = obj.get("LastModified")
                        if newest_time is None or (lm and lm > newest_time):
                            newest_time = lm
                            newest = key
                if newest:
                    gt_key = newest
                    if newest.lower().endswith(".pdf"):
                        break
            except Exception:
                continue

        now = datetime.utcnow().isoformat()
        conn = get_conn()
        try:
            uid = _ensure_user(conn, username, user_id, None, None)
            for it in (scan.get("items") or []):
                conn.execute(
                    "INSERT INTO outputs_history (user_id, case_id, label, ai_key, doctor_key, ground_truth_key, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (uid, case_id, it.get("label"), it.get("ai_key"), it.get("doctor_key"), gt_key, now),
                )
            if not (scan.get("items") or []) and gt_key:
                conn.execute(
                    "INSERT INTO outputs_history (user_id, case_id, label, ai_key, doctor_key, ground_truth_key, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (uid, case_id, None, None, None, gt_key, now),
                )
            conn.commit()
            return {"case_id": case_id, "count": conn.execute("SELECT COUNT(1) FROM outputs_history WHERE user_id=? AND case_id=?", (uid, case_id)).fetchone()[0]}
        finally:
            conn.close()


    @app.get("/history/cases")
    def history_cases(username: Optional[str] = None, user_id: Optional[int] = None) -> JSONResponse:
        conn = get_conn()
        try:
            uid = _ensure_user(conn, username, user_id, None, None)
            rows = conn.execute(
                "SELECT DISTINCT case_id FROM outputs_history WHERE user_id=? ORDER BY case_id DESC",
                (uid,),
            ).fetchall()
            cases = [r[0] for r in rows]
            return JSONResponse({"cases": cases})
        finally:
            conn.close()


    @app.get("/history/{case_id}")
    def history_list(case_id: str, username: Optional[str] = None, user_id: Optional[int] = None) -> Dict[str, Any]:
        conn = get_conn()
        try:
            uid = _ensure_user(conn, username, user_id, None, None)
            rows = conn.execute(
                "SELECT label, ai_key, doctor_key, ground_truth_key, created_at FROM outputs_history WHERE user_id=? AND case_id=? ORDER BY id DESC",
                (uid, case_id),
            ).fetchall()
            items = []
            for r in rows:
                items.append(
                    {
                        "label": r[0],
                        "ai_key": r[1],
                        "doctor_key": r[2],
                        "ground_truth_key": r[3],
                        "created_at": r[4],
                    }
                )
            return {"case_id": case_id, "items": items}
        finally:
            conn.close()


    ## (static route declared above to avoid shadowing by /history/{case_id})


    # n8n Workflow Integration Endpoints
    @app.post("/n8n/generate-report")
    async def generate_medical_report(patient_id: Optional[str] = None, username: Optional[str] = None, case_id: Optional[str] = None) -> Dict[str, Any]:
        """
        Triggers the complete medical report generation workflow via n8n.
        
        Args:
            patient_id: The patient ID to generate report for
            username: The username requesting the report
            
        Returns:
            Dict containing report generation status and details
        """
        try:
            # Normalize to case_id
            cid = (case_id or patient_id or "").strip()
            if not cid:
                raise HTTPException(status_code=400, detail="case_id required")
            # New: start main workflow via API/webhook and capture execution id
            try:
                capture = n8n_manager.trigger_main_workflow_and_capture_execution(cid, {"case_id": cid, "username": username})
            except Exception:
                capture = {"success": False}
            result = await report_generator.generate_complete_report(cid, username)
            if capture.get("success") and capture.get("execution_id"):
                # Persist captured execution id as well
                try:
                    conn = get_conn()
                    with conn:
                        conn.execute(
                            "INSERT INTO reports (case_id, status, started_at, metadata) VALUES (?, 'processing', datetime('now'), ?)",
                            (cid, json.dumps({"n8n_execution_id": capture["execution_id"], "source": "api_capture"})),
                        )
                except Exception:
                    pass
            
            # Store the report generation request in the database
            conn = get_conn()
            try:
                conn.execute(
                    """
                    INSERT INTO reports (case_id, email, status, started_at, metadata)
                    VALUES (?, ?, ?, ?, ?)
                    """,
                    (
                        cid,
                        username,
                        "processing",
                        datetime.utcnow().isoformat(),
                        json.dumps({
                            "source": "n8n_workflow",
                            "report_id": result.get("report_id"),
                            "workflow_data": result
                        }),
                    ),
                )
                conn.commit()
            finally:
                conn.close()
            
            return result
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@app.post("/webhook/progress")
def webhook_progress(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Endpoint for n8n to POST progress updates during workflow execution.
    Expected payload:
    {
        "case_id": "<user_case_id>",
        "progress": 45,
        "step": 2,
        "message": "Processing OCR data...",
        "timestamp": "2024-01-01T12:00:00Z"
    }
    """
    try:
        case_id = str(payload.get("case_id") or "").strip()
        progress = int(payload.get("progress", 0))
        step = int(payload.get("step", 0))
        message = str(payload.get("message", ""))
        timestamp = str(payload.get("timestamp", ""))
        
        if not case_id:
            raise ValueError("case_id required")
        
        # Store progress in database
        conn = get_conn()
        try:
            conn.execute(
                """
                INSERT INTO progress_updates (case_id, progress, step, message, timestamp)
                VALUES (?, ?, ?, ?, ?)
                """,
                (case_id, progress, step, message, timestamp),
            )
            conn.commit()
        finally:
            conn.close()
        
        return {"ok": True, "case_id": case_id, "progress": progress}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/webhook/finalize")
def webhook_finalize(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Endpoint for n8n to POST final artifacts/metrics when the workflow completes.
    Expected payload (examples accepted):
    {
        "case_id": "<user_case_id>",
        "pdf": {"signed_url": "...", "key": "..."},
        "docx": {"signed_url": "...", "key": "..."},
        "ocr_start_time": "...",
        "ocr_end_time": "...",
        "total_tokens_used": 123,
        "total_input_tokens": 456,
        "total_output_tokens": 789
    }
    Stores a row in runs and returns {ok: True}.
    """
    try:
        case_id = str(payload.get("case_id") or payload.get("patient_id") or "").strip()
        if not case_id:
            raise ValueError("case_id required")
        def _pick_url(obj: Any) -> str | None:
            if isinstance(obj, dict):
                return obj.get("signed_url") or obj.get("url") or obj.get("href")
            return None
        ai_url = _pick_url(payload.get("ai"))
        doc_url = _pick_url(payload.get("docx"))
        pdf_url = _pick_url(payload.get("pdf"))
        ocr_start = payload.get("ocr_start_time")
        ocr_end = payload.get("ocr_end_time")
        tot = payload.get("total_tokens_used")
        tin = payload.get("total_input_tokens")
        tout = payload.get("total_output_tokens")
        conn = get_conn()
        try:
            conn.execute(
                """
                INSERT INTO runs (case_id, created_at, ai_url, doc_url, pdf_url, ocr_start_time, ocr_end_time,
                                total_tokens_used, total_input_tokens, total_output_tokens)
                VALUES (?, datetime('now'), ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (case_id, ai_url, doc_url, pdf_url, ocr_start, ocr_end, tot, tin, tout),
            )
            conn.commit()
        finally:
            conn.close()
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


    @app.get("/runs/{case_id}")
    def get_runs(case_id: str) -> Dict[str, Any]:
        """Return latest run row for a case (for UI fallbacks)."""
        conn = get_conn()
        try:
            row = conn.execute(
                "SELECT * FROM runs WHERE case_id=? ORDER BY id DESC LIMIT 1",
                (case_id,),
            ).fetchone()
            if not row:
                return {"ok": True, "run": None}
            return {
                "ok": True,
                "run": {
                    "case_id": row[1],
                    "created_at": row[2],
                    "ai_url": row[3],
                    "doc_url": row[4],
                    "pdf_url": row[5],
                    "ocr_start_time": row[6],
                    "ocr_end_time": row[7],
                    "total_tokens_used": row[8],
                    "total_input_tokens": row[9],
                    "total_output_tokens": row[10],
                },
            }
        finally:
            conn.close()


@app.get("/runs/{case_id}/all")
def list_runs(case_id: str) -> Dict[str, Any]:
    """Return all runs for a case, newest first, so UI can map rows accurately."""
    conn = get_conn()
    try:
        rows = conn.execute(
            "SELECT * FROM runs WHERE case_id=? ORDER BY id DESC",
            (case_id,),
        ).fetchall()
        runs: list[Dict[str, Any]] = []
        for row in rows:
            runs.append(
                {
                    "case_id": row[1],
                    "created_at": row[2],
                    "ai_url": row[3],
                    "doc_url": row[4],
                    "pdf_url": row[5],
                    "ocr_start_time": row[6],
                    "ocr_end_time": row[7],
                    "total_tokens_used": row[8],
                    "total_input_tokens": row[9],
                    "total_output_tokens": row[10],
                }
            )
        return {"ok": True, "runs": runs}
    finally:
        conn.close()

@app.get("/progress/{case_id}/latest")
def get_latest_progress(case_id: str) -> Dict[str, Any]:
    """Get the latest progress update for a case."""
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT * FROM progress_updates WHERE case_id=? ORDER BY created_at DESC LIMIT 1",
            (case_id,),
        ).fetchone()
        if not row:
            return {"ok": True, "progress": None}
        return {
            "ok": True,
            "progress": {
                "case_id": row[1],
                "progress": row[2],
                "step": row[3],
                "message": row[4],
                "timestamp": row[5],
                "created_at": row[6],
            }
        }
    finally:
        conn.close()

@app.get("/progress/{case_id}/all")
def get_all_progress(case_id: str) -> Dict[str, Any]:
    """Get all progress updates for a case, newest first."""
    conn = get_conn()
    try:
        rows = conn.execute(
            "SELECT * FROM progress_updates WHERE case_id=? ORDER BY created_at DESC",
            (case_id,),
        ).fetchall()
        updates = []
        for row in rows:
            updates.append({
                "case_id": row[1],
                "progress": row[2],
                "step": row[3],
                "message": row[4],
                "timestamp": row[5],
                "created_at": row[6],
            })
        return {"ok": True, "updates": updates}
    finally:
        conn.close()

    @app.get("/n8n/report-status/{report_id}")
    def get_report_status(report_id: str) -> Dict[str, Any]:
        """
        Gets the status of a report generation process.
        
        Args:
            report_id: The report ID
            
        Returns:
            Dict containing report status
        """
        try:
            status = report_generator.get_report_status(report_id)
            return status
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to get report status: {str(e)}")


    @app.get("/n8n/workflow-status/{execution_id}")
    def get_workflow_status(execution_id: str) -> Dict[str, Any]:
        """
        Gets the status of a specific n8n workflow execution.
        
        Args:
            execution_id: The execution ID from the workflow
            
        Returns:
            Dict containing workflow status
        """
        try:
            status = n8n_manager.get_workflow_status(execution_id)
            return status
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to get workflow status: {str(e)}")


    @app.get("/n8n/workflow-result/{execution_id}")
    def get_workflow_result(execution_id: str) -> Dict[str, Any]:
        """
        Gets the result of a completed n8n workflow execution.
        
        Args:
            execution_id: The execution ID from the workflow
            
        Returns:
            Dict containing workflow results
        """
        try:
            result = n8n_manager.get_workflow_result(execution_id)
            return result
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to get workflow result: {str(e)}")


    @app.get("/n8n/download-report/{report_id}")
    def download_generated_report(report_id: str):
        """
        Downloads a generated medical report.
        
        Args:
            report_id: The report ID
            
        Returns:
            File response with the report
        """
        try:
            report_file = report_generator.get_report_file(report_id)
            if report_file and report_file.exists():
                return FileResponse(
                    path=str(report_file), 
                    filename=report_file.name,
                    media_type="application/pdf" if report_file.suffix == ".pdf" else "text/html"
                )
            else:
                raise HTTPException(status_code=404, detail="Report file not found")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to download report: {str(e)}")


    # Optional: admin cleanup to remove seeded demo data
    @app.delete("/admin/cleanup_seeds")
    def cleanup_seeds() -> Dict[str, int]:
        conn = get_conn()
        try:
            # Find seeded cycles
            rows = conn.execute("SELECT id FROM report_cycles WHERE metadata LIKE '%\"source\": \"seed\"%'").fetchall()
            ids = [int(r[0]) for r in rows]
            # Delete associated files
            for cid in ids:
                conn.execute("DELETE FROM report_files WHERE cycle_id=?", (cid,))
            conn.execute("DELETE FROM report_cycles WHERE id IN (" + ",".join([str(i) for i in ids]) + ")") if ids else None
            conn.commit()
            return {"deleted_cycles": len(ids)}
        finally:
            conn.close()


    def _convert_docx_to_pdf_local(docx_path: str, out_pdf_path: str) -> bool:
        """Try to convert DOCX to PDF locally using docx2pdf or LibreOffice.
        Returns True if out_pdf_path was created with content.
        """
        try:
            from docx2pdf import convert as _docx2pdf
        except Exception:
            _docx2pdf = None
        from pathlib import Path as _Path
        import subprocess, shutil

        pdf_path = _Path(out_pdf_path)
        # Try docx2pdf first
        if _docx2pdf is not None:
            try:
                _docx2pdf(docx_path, out_pdf_path)
                if pdf_path.exists() and pdf_path.stat().st_size > 0:
                    return True
            except Exception:
                pass
        # Try LibreOffice fallback
        try:
            soffice = shutil.which("soffice") or shutil.which("soffice.exe")
            if soffice:
                tmpdir = str(_Path(out_pdf_path).parent)
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                    check=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                lo_pdf = _Path(docx_path).with_suffix(".pdf")
                if lo_pdf.exists() and lo_pdf.stat().st_size > 0:
                    lo_pdf.rename(pdf_path)
                    return True
        except Exception:
            pass
        return False
    